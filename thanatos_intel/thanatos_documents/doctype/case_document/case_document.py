"""Case Document — sistema documentale e di protocollo.

Ciclo: Bozza (docx editabile) -> Certificato (PDF + SHA-256, bloccato) ->
Inviato (protocollo) -> Archiviato. Ogni ri-certificazione crea una nuova versione.
"""
import hashlib
import os
import subprocess
import tempfile

import frappe
from frappe.model.document import Document
from frappe.utils import now_datetime, nowdate, getdate

SOFFICE = next((p for p in ("/usr/bin/soffice", "/usr/bin/libreoffice") if os.path.exists(p)), None)


def _docx_to_pdf(docx_path):
    if not SOFFICE:
        frappe.throw("LibreOffice non installato sul server.")
    outdir = tempfile.mkdtemp()
    env = dict(os.environ, HOME=outdir)
    r = subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                       capture_output=True, text=True, timeout=180, env=env)
    pdfs = [x for x in os.listdir(outdir) if x.endswith(".pdf")]
    if not pdfs:
        frappe.throw("Conversione PDF fallita: " + (r.stderr or r.stdout)[:200])
    return os.path.join(outdir, pdfs[0])


def _attach(doc, fname, content, field=None):
    f = frappe.get_doc({"doctype": "File", "file_name": fname, "is_private": 1, "content": content,
                        "attached_to_doctype": "Case Document", "attached_to_name": doc.name,
                        "attached_to_field": field})
    f.save(ignore_permissions=True)
    return f.file_url


def _site_path(file_url):
    private = "/private/files/" in file_url
    return frappe.get_site_path("private" if private else "public", "files", file_url.split("/files/")[-1])


class CaseDocument(Document):

    def before_insert(self):
        if self.investigation_case and not self.client:
            self.client = frappe.db.get_value("Investigation Case", self.investigation_case, "client")

    # ----- generazione bozza docx dai report del case -----
    @frappe.whitelist()
    def generate_from_reports(self, kinds):
        from docx import Document as Docx
        from docx.shared import Pt, RGBColor
        from thanatos_intel.reporting.case_reports import _sections, _strip, _client_name, KIND_LABEL

        if isinstance(kinds, str):
            kinds = [k.strip() for k in kinds.split(",") if k.strip()]
        case = frappe.get_doc("Investigation Case", self.investigation_case)
        client = _client_name(case)

        d = Docx()
        st = d.styles["Normal"]
        st.font.name = "Calibri"; st.font.size = Pt(10.5)
        t = d.add_heading(self.title or "Documento", level=0)
        d.add_paragraph("Cliente: %s  ·  %s  ·  %s" % (client, self.investigation_case, nowdate()))
        seen = set()
        for kind in kinds:
            for stitle, body in _sections(case, kind):
                if stitle in seen:
                    continue
                seen.add(stitle)
                d.add_heading(stitle, level=1)
                for para in _strip(body).split("\n\n"):
                    para = para.strip()
                    if para:
                        d.add_paragraph(para)
        d.add_paragraph("")
        foot = d.add_paragraph("Documento di lavoro Thanatos Intel — modificabile fino alla certificazione.")
        foot.runs[0].italic = True

        buf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        d.save(buf.name)
        content = open(buf.name, "rb").read()
        os.unlink(buf.name)
        fname = "%s_v%s_bozza.docx" % (self.name, self.version or 1)
        url = _attach(self, fname, content, "editable_docx")
        self.db_set("editable_docx", url)
        self.db_set("status", "Bozza")
        frappe.db.commit()
        return {"editable_docx": url, "sections": len(seen)}

    # ----- certificazione: docx -> pdf + sha256, blocco versione -----
    @frappe.whitelist()
    def certify(self):
        if not self.editable_docx:
            frappe.throw("Allegare/generare prima il docx editabile.")
        src = _site_path(self.editable_docx)
        if not os.path.exists(src):
            frappe.throw("Docx non trovato sul filesystem.")
        pdf_path = _docx_to_pdf(src)
        content = open(pdf_path, "rb").read()
        sha = hashlib.sha256(content).hexdigest()
        fname = "%s_v%s_certificato.pdf" % (self.name, self.version or 1)
        url = _attach(self, fname, content, "certified_pdf")

        self.append("versions", {"version": self.version or 1, "sha256": sha,
                                 "certified_pdf": url, "certified_on": now_datetime(),
                                 "certified_by": frappe.session.user, "note": "certificazione"})
        self.certified_pdf = url
        self.sha256 = sha
        self.certified_on = now_datetime()
        self.certified_by = frappe.session.user
        self.status = "Certificato"
        self.save(ignore_permissions=True)

        # push in Drive con tag cliente
        try:
            from thanatos_intel.reporting.case_reports import _put_in_drive, _client_name
            client = _client_name(frappe.get_doc("Investigation Case", self.investigation_case))
            self.drive_file = _put_in_drive(self.investigation_case, fname, content, "application/pdf", client) or self.drive_file
            self.db_set("drive_file", self.drive_file)
        except Exception:
            frappe.log_error(frappe.get_traceback(), "case_document drive")
        frappe.db.commit()
        return {"certified_pdf": url, "sha256": sha, "version": self.version}

    # ----- nuova versione (ri-apre alla modifica) -----
    @frappe.whitelist()
    def new_version(self):
        if self.status not in ("Certificato", "Inviato", "Archiviato"):
            frappe.throw("Si crea una nuova versione solo da un documento gia certificato.")
        self.version = (self.version or 1) + 1
        self.status = "Bozza"
        self.certified_pdf = None
        self.sha256 = None
        self.certified_on = None
        self.certified_by = None
        self.save(ignore_permissions=True)
        frappe.db.commit()
        return {"version": self.version}

    # ----- protocollo -----
    @frappe.whitelist()
    def assign_protocol(self, direction=None):
        if self.protocol_number:
            return {"protocol_number": self.protocol_number}
        if direction:
            self.direction = direction
        year = getdate().year
        prefix = "PROT-%s/" % year
        last = frappe.db.sql(
            "select protocol_number from `tabCase Document` where protocol_number like %s order by protocol_number desc limit 1",
            (prefix + "%",))
        nxt = 1
        if last and last[0][0]:
            try:
                nxt = int(last[0][0].split("/")[-1]) + 1
            except Exception:
                nxt = 1
        self.protocol_number = "%s%05d" % (prefix, nxt)
        self.protocol_date = nowdate()
        self.save(ignore_permissions=True)
        frappe.db.commit()
        return {"protocol_number": self.protocol_number, "protocol_date": self.protocol_date}

    # ----- invio -----
    @frappe.whitelist()
    def mark_sent(self, channel, recipient=None):
        if self.status not in ("Certificato",):
            frappe.throw("Inviare solo documenti certificati.")
        if not self.protocol_number:
            self.assign_protocol("Uscita")
            self.reload()
        self.sent_channel = channel
        self.sent_to = recipient
        self.sent_on = now_datetime()
        self.status = "Inviato"
        self.save(ignore_permissions=True)
        frappe.db.commit()
        return {"status": "Inviato", "protocol_number": self.protocol_number}

    @frappe.whitelist()
    def send_docuseal(self, signer_email=None, signer_name=None):
        if not self.certified_pdf:
            frappe.throw("Certificare prima il documento.")
        from thanatos_intel.reporting.case_reports import send_report_to_docuseal
        r = send_report_to_docuseal(self.certified_pdf, self.investigation_case, signer_email, signer_name)
        if r.get("ok"):
            self.mark_sent("DocuSeal", signer_email)
        return r
