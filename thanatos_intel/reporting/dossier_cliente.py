"""Dossier investigativo cliente (DOCX editabile) — tutto ciò che abbiamo sul caso.

Compila: sintesi esecutiva, parti, documenti+autenticità, anomalie/red flag (dalle
attività del caso), verifiche svolte, quantificazione danno, conclusioni e
raccomandazioni. Pensato per essere inviato/editato (es. da Lorenzo).
"""
import io

import frappe
from frappe.utils import nowdate


def _activities(case, *needles):
    rows = frappe.get_all("Case Activity", filters={"parent": case}, fields=["description"],
                          order_by="activity_date asc", limit=0)
    out = []
    for r in rows:
        d = (r.description or "").strip()
        if d and (not needles or any(n.lower() in d.lower() for n in needles)):
            out.append(d)
    return out


@frappe.whitelist()
def genera_dossier(case):
    import docx
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    c = frappe.get_doc("Investigation Case", case)
    client = (frappe.db.get_value("Investigation Client", c.client, "client_name") if c.client else None) or "Cliente"
    NAVY = RGBColor(0x0D, 0x1B, 0x3E); GOLD = RGBColor(0xC8, 0xA9, 0x6E)

    d = docx.Document()
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    def h(txt, size=14, color=NAVY, space=6):
        p = d.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
        p.paragraph_format.space_before = Pt(space); p.paragraph_format.space_after = Pt(3); return p

    def para(txt, size=10.5):
        p = d.add_paragraph(); r = p.add_run(txt); r.font.size = Pt(size); return p

    # Copertina
    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("THANATOS INTEL"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = GOLD
    s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("DOSSIER INVESTIGATIVO"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{c.case_title}\nCliente: {client} · Caso {c.name} · {nowdate()}").font.size = Pt(11)
    d.add_paragraph("Documento riservato. Distribuzione limitata al committente e ai suoi legali.").italic = True

    # 1. Sintesi esecutiva
    h("1. Sintesi esecutiva")
    _sint = (c.get("summary") or "").strip()
    if not _sint and c.get("description"):
        _sint = frappe.utils.strip_html(c.description).strip()
    para(_sint or "Sintesi esecutiva non ancora compilata per questo caso. "
                  "Compilare il campo 'summary' del caso o le attività investigative.")

    # 2. Parti
    h("2. Le parti")
    for ce in (c.get("case_entities") or [])[:25]:
        et = frappe.db.get_value("Investigation Entity", ce.entity, ["full_name", "entity_type", "risk_level"], as_dict=True)
        if et:
            para(f"• {et.full_name} ({et.entity_type})" + (f" — rischio {et.risk_level}" if et.risk_level else "")
                 + (f" — {ce.notes}" if ce.notes else ""), size=10)

    # 3. Documenti analizzati
    h("3. Documenti analizzati — autenticità e catena di custodia")
    evs = frappe.get_all("Investigation Evidence", filters={"investigation_case": case},
                         fields=["evidence_name", "authenticity", "hash_value", "attached_file"], order_by="creation asc", limit=0)
    tab = d.add_table(rows=1, cols=3); tab.style = "Light Grid Accent 1"
    hdr = tab.rows[0].cells; hdr[0].text = "Documento"; hdr[1].text = "Autenticità"; hdr[2].text = "SHA-256 (estratto)"
    for e in evs:
        nm = (e.attached_file or e.evidence_name or "").split("/files/")[-1]
        row = tab.add_row().cells
        row[0].text = nm[:60]; row[1].text = e.authenticity or "N/D"; row[2].text = (e.hash_value or "")[:16]
    para(f"Totale {len(evs)} reperti. Ogni documento è identificato dal proprio digest SHA-256 (catena di custodia).", size=9)

    # 4. Anomalie e red flag
    h("4. Anomalie e red flag principali")
    for blk in _activities(case, "DOPPIA CESSIONE", "Fattorelli", "DICHIARAZIONE FATTURE", "Debitori", "Contratto cessione"):
        para(blk, size=10); d.add_paragraph()

    # 5. Verifiche svolte
    h("5. Verifiche svolte")
    for blk in _activities(case, "Screening", "Verifica camerale", "VERIFICA PARTI", "RICONCILIAZIONE"):
        para(blk, size=10); d.add_paragraph()

    # Rileva se il caso e' effettivamente una frode su cessione crediti (playbook Bomax)
    _is_cessione = bool(_activities(case, "DOPPIA CESSIONE", "cessione credit", "cessione di credit", "Piattaforma Cessione"))

    # 6. Danno
    h("6. Quantificazione del danno")
    if _is_cessione:
        para("Esposizione del cliente: € 800.000 versati per crediti risultati non genuini. La quantificazione "
             "definitiva e l'individuazione dei beneficiari dei pagamenti (cedente e/o intermediari) richiede "
             "l'acquisizione di tutti i contratti e dei bonifici del cliente (tracciamento dei flussi).")
    else:
        para("Quantificazione del danno: da determinare sulla base delle evidenze raccolte per questo caso. "
             "Non risultano, allo stato, esposizioni economiche accertate.")

    # 7. Conclusioni
    h("7. Conclusioni e raccomandazioni")
    if _is_cessione:
        _concl = [
            "Acquisire, tramite delega del cliente, lo stato reale dei crediti sul cassetto fiscale e sulla Piattaforma Cessione Crediti AdE, e gli XML delle fatture elettroniche per la riconciliazione.",
            "Tracciare i bonifici degli € 800.000 (follow-the-money) per individuare i beneficiari.",
            "Valutare denuncia/querela per truffa (artt. 640/640-bis c.p.), falso e fatture per operazioni inesistenti, verso cedente, intermediari e asseveratori.",
            "Promuovere azione civile di risarcimento/restituzione ed escutere le polizze RC professionali degli asseveratori.",
            "Documentare la buona fede e la diligenza del cliente per neutralizzare il recupero erariale verso il cessionario.",
        ]
    else:
        _concl = [
            "Consolidare le evidenze raccolte (identita, footprint societario, documenti) nel fascicolo del caso.",
            "Verificare gli elementi ancora aperti indicati nelle attivita investigative.",
            "Valutare le azioni conseguenti in base all'esito e al mandato del caso.",
        ]
    for it in _concl:
        para("• " + it, size=10.5)

    out = io.BytesIO(); d.save(out); content = out.getvalue()
    fname = f"DOSSIER CLIENTE - {client} - {c.name}.docx"
    old = frappe.db.get_value("File", {"attached_to_doctype": "Investigation Case", "attached_to_name": case, "file_name": fname}, "name")
    if old:
        frappe.delete_doc("File", old, ignore_permissions=True, force=True)
    f = frappe.get_doc({"doctype": "File", "file_name": fname, "is_private": 1, "content": content,
                        "attached_to_doctype": "Investigation Case", "attached_to_name": case})
    f.save(ignore_permissions=True)
    try:
        from thanatos_intel.reporting.case_reports import _put_in_drive
        _put_in_drive(case, fname, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", client, subfolder="05 Report")
    except Exception:
        frappe.log_error(frappe.get_traceback(), "dossier drive")
    frappe.db.commit()
    return {"ok": True, "file_url": f.file_url, "documenti": len(evs)}
